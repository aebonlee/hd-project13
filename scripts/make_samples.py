# -*- coding: utf-8 -*-
"""
샘플 사양서 3종 생성 — 실행: python3 scripts/make_samples.py

**업체마다 양식이 다른 상황을 일부러 재현한다.**
셋이 똑같으면 추출기가 한 가지만 맞춰도 통과해 버려 검사가 의미를 잃는다.

  · 엑셀 (A사) — ID 만 있음. PGN·SA 는 ID 에서 계산해야 한다.
  · 워드 (B사) — PGN·SA·채널만 있음. ID 는 없다 → 지어내면 안 된다.
  · PDF  (C사) — 다섯 항목이 다 있음. 대신 열 이름이 영문이고 한 줄이 어긋나 있다.
"""
import os, sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, 'sample')
os.makedirs(OUT, exist_ok=True)

# ── A사: 엑셀 — ID 만 ──────────────────────────────────────────────────
def make_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook(); ws = wb.active; ws.title = 'CAN 사양'
    ws['A1'] = 'HX220A 전자제어 사양서 (A사)'
    ws['A1'].font = Font(bold=True, size=13)
    ws['A2'] = '작성일'; ws['B2'] = '2026-08-20'
    ws['A3'] = '전장품'; ws['B3'] = 'ECM (엔진제어기)'
    head = ['메시지명', 'ID', 'Period', '길이(byte)', 'SPN', '신호명', '시작비트', '비트길이', '분해능', '단위']
    ws.append([]); ws.append(head)
    hr = ws.max_row
    fill = PatternFill('solid', fgColor='1F4E79'); font = Font(color='FFFFFF', bold=True, size=10)
    for c in range(1, len(head) + 1):
        cell = ws.cell(row=hr, column=c); cell.fill = fill; cell.font = font
        cell.alignment = Alignment(horizontal='center')
    rows = [
        ['EEC1', '0x0CF00400', '20ms',  8, 190, '엔진 회전수',     25, 16, 0.125, 'rpm'],
        ['EEC1', '0x0CF00400', '20ms',  8, 513, '실제 토크비',      3,  8, 1,     '%'],
        ['ET1',  '0x18FEEE00', '1000ms',8, 110, '냉각수 온도',      1,  8, 1,     '℃'],
        ['LFE',  '0x18FEF200', '100ms', 8, 183, '연료 소비율',      1, 16, 0.05,  'L/h'],
        ['CCVS', '0x18FEF100', '100ms', 8,  84, '차속',            2, 16, 1/256, 'km/h'],
    ]
    for r in rows: ws.append(r)
    for i, w in enumerate([12, 14, 10, 12, 8, 18, 10, 10, 10, 8], start=1):
        ws.column_dimensions[chr(64 + i)].width = w
    p = os.path.join(OUT, 'A사_ECM_CAN사양_20260820.xlsx')
    wb.save(p); return p

# ── B사: 워드 — PGN·SA·채널만 ──────────────────────────────────────────
def make_docx():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading('B사 계기판(Cluster) CAN 사양서', level=1)
    doc.add_paragraph('모델: HX220A / HX300L    전장품: Cluster    개정: Rev.3')
    doc.add_paragraph('※ 본 사양서는 PGN 기준으로 작성되었습니다. (ID 는 차량 통합 시 확정)')

    head = ['메시지명', 'PGN', 'CAN Channel', 'Source Address', '주기', '신호명', '단위']
    rows = [
        ['DASH',  'FEFC', '1', '17', '1000ms', '연료 잔량',     '%'],
        ['DASH',  'FEFC', '1', '17', '1000ms', '요소수 잔량',   '%'],
        ['AMB',   'FEF5', '2', '17', '1000ms', '외기 온도',     '℃'],
        ['HOURS', 'FEE5', '1', '17', 'On Request', '엔진 가동시간', 'h'],
        ['VEP1',  'FEF7', '2', '17', '1000ms', '배터리 전압',   'V'],
    ]
    t = doc.add_table(rows=1, cols=len(head)); t.style = 'Table Grid'
    for i, h in enumerate(head):
        cell = t.rows[0].cells[i]; cell.text = h
        for r in cell.paragraphs[0].runs: r.font.bold = True; r.font.size = Pt(9)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)
            for run in cells[i].paragraphs[0].runs: run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph('부속 표 (CAN 사양 아님 — 추출 대상에서 걸러져야 함)')
    t2 = doc.add_table(rows=1, cols=3); t2.style = 'Table Grid'
    for i, h in enumerate(['항목', '값', '비고']): t2.rows[0].cells[i].text = h
    for r in [['공급전압', 'DC 24V', ''], ['동작온도', '-30 ~ 80℃', '']]:
        cells = t2.add_row().cells
        for i, v in enumerate(r): cells[i].text = v

    p = os.path.join(OUT, 'B사_Cluster_CAN사양_Rev3.docx')
    doc.save(p); return p

# ── C사: PDF — 영문 머리글 · 한 줄이 어긋남 ────────────────────────────
def make_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    KO = 'HYSMyeongJo-Medium'
    pdfmetrics.registerFont(UnicodeCIDFont(KO))
    p = os.path.join(OUT, 'C사_BodyECU_CAN_Spec.pdf')
    W, H = A4
    c = canvas.Canvas(p, pagesize=A4)
    c.setFont('Helvetica-Bold', 15)
    c.drawString(20 * mm, H - 25 * mm, 'C Corp. Body ECU - J1939 CAN Specification')
    c.setFont(KO, 10)
    c.drawString(20 * mm, H - 32 * mm, '모델: HW250 / 전장품: Body ECU / Rev. 1.2')
    c.line(20 * mm, H - 35 * mm, W - 20 * mm, H - 35 * mm)

    head = ['Message', 'Identifier', 'PGN', 'CAN Ch', 'Src Addr', 'Cycle Time']
    rows = [
        ['LIGHT',  '0x18FE4021', 'FE40', '1', '21', '500 ms'],
        ['DOOR',   '0x18FE4121', 'FE41', '1', '21', '500 ms'],
        # ⚠ 아래 줄은 일부러 PGN 이 ID 와 어긋나 있다 (사람이 손으로 옮기다 낸 실수 재현)
        ['WIPER',  '0x18FE4221', 'FE43', '2', '21', '1000 ms'],
        ['HVAC',   '0x18FEF421', 'FEF4', '3', '21', '100 ms'],
    ]
    xs = [22 * mm, 48 * mm, 84 * mm, 106 * mm, 126 * mm, 152 * mm]
    y = H - 48 * mm
    c.setFont('Helvetica-Bold', 9)
    for x, h in zip(xs, head): c.drawString(x, y, h)
    c.line(20 * mm, y - 2 * mm, W - 20 * mm, y - 2 * mm)
    c.setFont('Helvetica', 9)
    y -= 8 * mm
    for r in rows:
        for x, v in zip(xs, r): c.drawString(x, y, v)
        y -= 6.5 * mm
    c.setFont(KO, 8)
    c.drawString(20 * mm, y - 6 * mm, '※ WIPER 행은 PGN 표기가 ID 와 다릅니다 (검증용).')
    c.showPage(); c.save(); return p

made = []
for fn in (make_xlsx, make_docx, make_pdf):
    try: made.append(fn())
    except Exception as e: print('  실패:', fn.__name__, e); sys.exit(1)
for m in made: print('  ' + os.path.relpath(m, BASE) + '  ' + str(os.path.getsize(m)) + ' bytes')
